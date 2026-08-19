import os
import time
import re
import threading
import io
import requests
from typing import Callable, Optional, Set
from playwright.sync_api import sync_playwright, Page, BrowserContext, Browser

# Anti-blur JS payload to override visibilityState and blur events
ANTI_BLUR_SCRIPT = """
Object.defineProperty(document, 'hidden', { get: () => false, configurable: true });
Object.defineProperty(document, 'visibilityState', { get: () => 'visible', configurable: true });
window.addEventListener('visibilitychange', (e) => e.stopImmediatePropagation(), true);
window.addEventListener('blur', (e) => e.stopImmediatePropagation(), true);
window.onblur = null;
document.onvisibilitychange = null;
"""

class ADBAutomationEngine:
    """ADB Online Eğitim Otomasyon ve Word Dökümantasyon Motoru (Playwright Tabanlı)"""

    def __init__(self, log_callback: Optional[Callable[[str], None]] = None, status_callback: Optional[Callable[[str], None]] = None):
        self.log_callback = log_callback or print
        self.status_callback = status_callback or (lambda s: None)
        
        self.playwright = None
        self.browser: Optional[Browser] = None
        self.context: Optional[BrowserContext] = None
        self.page: Optional[Page] = None
        
        self.is_running = False
        self.is_paused = False
        self.playback_speed = 2.0
        self.auto_next = True
        self.mute_audio = True
        self.save_word_docs = True
        
        self.thread: Optional[threading.Thread] = None
        self.visited_pages: Set[str] = set()
        self.output_dir = os.path.join(os.getcwd(), "Egitim_Dokumanlari")

    def log(self, message: str):
        self.log_callback(message)

    def set_status(self, status: str):
        self.status_callback(status)

    def launch_browser(self):
        """Kullanıcının giriş yapabilmesi için gerçek tarayıcı penceresini açar."""
        try:
            self.log("🚀 Tarayıcı başlatılıyor (Chromium)...")
            self.playwright = sync_playwright().start()
            self.browser = self.playwright.chromium.launch(
                headless=False,
                args=[
                    "--start-maximized",
                    "--disable-blink-features=AutomationControlled",
                    "--no-sandbox"
                ]
            )
            self.context = self.browser.new_context(no_viewport=True)
            self.page = self.context.new_page()

            # Anti-blur injection
            self.page.add_init_script(ANTI_BLUR_SCRIPT)

            self.log("🌐 https://adbs.uab.gov.tr/ adresine gidiliyor...")
            self.page.goto("https://adbs.uab.gov.tr/users/my-educations", wait_until="domcontentloaded")
            self.set_status("Giriş Bekleniyor")
            self.log("🔑 Lütfen açılan pencerede e-Devlet ile giriş yapınız. Giriş yaptıktan sonra '▶️ 2. Otomasyonu Başlat' butonuna tıklayınız.")
        except Exception as e:
            self.log(f"❌ Tarayıcı başlatma hatası: {str(e)}")
            self.stop()

    def start_automation_loop(self):
        """Otomasyon döngüsünü arka planda başlatır."""
        if not self.page:
            self.log("⚠️ Önce tarayıcıyı başlatmalısınız!")
            return
        
        self.is_running = True
        self.is_paused = False
        self.thread = threading.Thread(target=self._run_loop, daemon=True)
        self.thread.start()
        self.log("▶️ Otomasyon döngüsü başlatıldı.")
        self.set_status("Çalışıyor")

    def pause(self):
        self.is_paused = True
        self.log("⏸️ Otomasyon duraklatıldı.")
        self.set_status("Duraklatıldı")

    def resume(self):
        self.is_paused = False
        self.log("▶️ Otomasyon devam ettiriliyor.")
        self.set_status("Çalışıyor")

    def stop(self):
        self.is_running = False
        self.log("🛑 Otomasyon durduruluyor...")
        try:
            if self.browser:
                self.browser.close()
            if self.playwright:
                self.playwright.stop()
        except Exception as e:
            pass
        self.set_status("Durduruldu")
        self.log("✅ Tarayıcı kapatıldı.")

    def _run_loop(self):
        """Ana otomasyon döngüsü."""
        last_click_time = 0

        while self.is_running:
            try:
                if self.is_paused or not self.page:
                    time.sleep(1)
                    continue

                current_url = self.page.url

                # 1. Anti-blur kontrolü yap
                try:
                    self.page.evaluate(ANTI_BLUR_SCRIPT)
                except Exception:
                    pass

                # 2. Sayfa içeriğini Word Dökümanı olarak kaydet (Yazı, Başlık, Tablo, Resimler)
                if self.save_word_docs and "my-educations" not in current_url and "giris" not in current_url:
                    self._capture_page_to_word()

                # 3. Videoları denetle ve oynat
                try:
                    video_count = self.page.locator("video").count()
                    if video_count > 0:
                        for i in range(video_count):
                            video = self.page.locator("video").nth(i)
                            
                            # Mute & speed set
                            if self.mute_audio:
                                self.page.evaluate("(v) => { v.muted = true; }", video.element_handle())
                            
                            self.page.evaluate(f"(v) => {{ v.playbackRate = {self.playback_speed}; }}", video.element_handle())

                            # Play if paused
                            is_paused = self.page.evaluate("(v) => v.paused && !v.ended", video.element_handle())
                            if is_paused:
                                self.page.evaluate("(v) => v.play()", video.element_handle())
                                self.log(f"▶️ Video {i+1} oynatılıyor ({self.playback_speed}x hızında).")

                            # Check if ended
                            is_ended = self.page.evaluate("(v) => v.ended", video.element_handle())
                            if is_ended and self.auto_next:
                                self.log(f"✅ Video {i+1} tamamlandı. Sonraki adıma geçiliyor...")
                                self._trigger_next_button()
                except Exception:
                    pass

                # 4. Sayaç / Geri Sayım ("Kalan Zaman: 00:06") ve "İleri / Devam Et" Butonu Kontrolü
                now = time.time()
                if self.auto_next and (now - last_click_time > 2.0):
                    self._check_timer_counter()

                    clicked = self._trigger_next_button()
                    if clicked:
                        last_click_time = now

                # 5. Kurs listesi sayfasında ise otomatik derse gir
                if "my-educations" in current_url:
                    self._check_education_list()

            except Exception as e:
                time.sleep(1)

            time.sleep(1.5)

    def _check_timer_counter(self):
        """Sayfadaki geri sayım sayacını ('Kalan Zaman: 00:06') kontrol eder."""
        try:
            page_text = self.page.locator("body").inner_text()
            matches = re.findall(r'(?:kalan zaman|kalan süre|süre|sayaç)?\s*:?\s*(\d{1,2}:\d{2}|\d+\s*sn|\d+\s*saniye)', page_text, re.IGNORECASE)
            if matches:
                timer_str = matches[0]
                if not hasattr(self, '_last_timer_logged') or time.time() - getattr(self, '_last_timer_logged', 0) > 4.0:
                    self.log(f"⏳ Kalan Zaman: {timer_str} (Süre bitince otomatik 'İleri' butonuna basılacak)...")
                    self._last_timer_logged = time.time()
        except Exception:
            pass

    def _trigger_next_button(self) -> bool:
        """Ekranda aktif olan 'İleri / Devam Et / Sonraki Ders' butonlarını bulur ve tıklar."""
        if not self.page:
            return False

        target_texts = [
            'ileri', '> ileri', 'ileri >', '>ileri',
            'devam et', 'devam', 'sonraki ders', 'sonraki konu', 
            'sonraki', 'eğitimi tamamla', 'eğitime başla', 
            'tamam', 'ok', 'dersi bitir', 'eğitime devam et'
        ]

        try:
            buttons = self.page.locator("button, a, input[type='button'], input[type='submit'], .btn").all()
            for btn in buttons:
                if not btn.is_visible():
                    continue

                raw_text = (btn.inner_text() or btn.get_attribute("value") or "").strip()
                clean_text = raw_text.lower().replace(">", "").strip()

                if any(t == clean_text or t in clean_text for t in ['ileri', 'devam', 'sonraki', 'tamamla', 'başla', 'tamam', 'ok', 'bitir']):
                    # Disabled kontrolü
                    is_disabled = self.page.evaluate("""(el) => {
                        return el.disabled || el.classList.contains('disabled') || getComputedStyle(el).pointerEvents === 'none';
                    }""", btn.element_handle())

                    if is_disabled:
                        # Geri sayım henüz bitmedi (buton kilitli)
                        continue

                    self.log(f"👉 Otomatik İlerleme: '{raw_text.upper()}' butonuna tıklandı!")
                    btn.click()
                    return True

            # Modal onay dialogları
            modals = self.page.locator(".modal button, .swal2-confirm, .ngx-modal button").all()
            for mbtn in modals:
                if mbtn.is_visible():
                    is_disabled = self.page.evaluate("(el) => el.disabled || el.classList.contains('disabled')", mbtn.element_handle())
                    if not is_disabled:
                        self.log(f"👉 Modal onay butonuna tıklandı: '{mbtn.inner_text()}'")
                        mbtn.click()
                        return True
        except Exception:
            pass

        return False

    def _check_education_list(self):
        """Kayıtlı eğitimler listesinde başlanmamış/tamamlanmamış eğitime tıklar."""
        try:
            btns = self.page.locator("button, a").all()
            for btn in btns:
                if not btn.is_visible():
                    continue
                txt = (btn.inner_text() or "").strip().lower()
                if "eğitime başla" in txt or "eğitime devam et" in txt:
                    self.log(f"📚 Bulunan eğitime başlanıyor: '{txt}'")
                    btn.click()
                    time.sleep(2)
                    break
        except Exception:
            pass

    def _capture_page_to_word(self):
        """Ekrandaki Ders Başlığı, Metinler, Tablolar (Fonotik Alfabe vb.) ve Resimleri Word (.docx) olarak kaydeder."""
        try:
            current_url = self.page.url
            page_title = self.page.title() or "ADB Eğitim Notları"

            # Sayfa unik kimliği (tekrar kaydetmemek için)
            body_text = self.page.evaluate('() => document.body.innerText') or ""
            page_id = f"{current_url}_{len(body_text)}"
            if page_id in self.visited_pages:
                return

            self.visited_pages.add(page_id)

            from docx import Document
            from docx.shared import Inches, Pt, RGBColor

            os.makedirs(self.output_dir, exist_ok=True)

            clean_title = re.sub(r'[\\/*?:"<>|]', '', page_title).strip() or "ADB_Egitim_Notlari"
            docx_filename = os.path.join(self.output_dir, f"{clean_title}.docx")

            if os.path.exists(docx_filename):
                doc = Document(docx_filename)
            else:
                doc = Document()
                heading = doc.add_heading(page_title, level=0)
                heading.style.font.color.rgb = RGBColor(30, 58, 138)

            # 1. Ana Başlıklar (Örn: "FONOTİK ALFABE")
            sub_headings = self.page.locator("h1, h2, h3, h4, .card-header, .lesson-title, header").all()
            for h in sub_headings:
                if h.is_visible():
                    heading_text = (h.inner_text() or "").strip()
                    if heading_text and len(heading_text) < 100:
                        doc.add_heading(heading_text, level=1)
                        break

            # 2. Sayfadaki Tablolar (Örn: Fonotik Alfabe Tablosu)
            tables = self.page.locator("table").all()
            added_table_count = 0
            for tbl in tables:
                try:
                    if not tbl.is_visible():
                        continue
                    rows = tbl.locator("tr").all()
                    if len(rows) > 0:
                        doc_table = doc.add_table(rows=0, cols=0)
                        doc_table.style = 'Table Grid'
                        for row in rows:
                            cells = row.locator("th, td").all()
                            if len(cells) > 0:
                                # Kolon sayısını ayarla
                                while len(doc_table.columns) < len(cells):
                                    doc_table.add_column(Inches(1.5))
                                row_cells = doc_table.add_row().cells
                                for c_idx, cell in enumerate(cells):
                                    if c_idx < len(row_cells):
                                        row_cells[c_idx].text = (cell.inner_text() or "").strip()
                        added_table_count += 1
                        doc.add_paragraph() # Boş satır
                except Exception:
                    pass

            # 3. Metin Paragrafları
            paragraphs = self.page.locator("p, article, .content, .description, .lesson-text").all()
            added_text_count = 0
            for p in paragraphs:
                if p.is_visible():
                    txt = (p.inner_text() or "").strip()
                    if len(txt) > 10 and not txt.startswith("Kalan Zaman"):
                        p_elem = doc.add_paragraph(txt)
                        p_elem.style.font.name = 'Calibri'
                        p_elem.style.font.size = Pt(11)
                        added_text_count += 1

            # 4. Görseller / Resimler
            images = self.page.locator("img").all()
            added_img_count = 0
            for img in images:
                try:
                    if not img.is_visible():
                        continue
                    src = img.get_attribute("src")
                    if not src or "icon" in src.lower() or "logo" in src.lower() or "avatar" in src.lower():
                        continue

                    img_bytes = None
                    if src.startswith("data:image"):
                        import base64
                        base64_data = src.split(",")[1]
                        img_bytes = base64.b64decode(base64_data)
                    else:
                        if src.startswith("/"):
                            src = "https://adbs.uab.gov.tr" + src
                        res = requests.get(src, timeout=5)
                        if res.status_code == 200:
                            img_bytes = res.content

                    if img_bytes:
                        image_stream = io.BytesIO(img_bytes)
                        doc.add_picture(image_stream, width=Inches(5.5))
                        added_img_count += 1
                except Exception:
                    pass

            doc.save(docx_filename)
            self.log(f"📄 Sayfa Word belgesine aktarıldı: '{clean_title}.docx' (Tablo: {added_table_count}, Paragraf: {added_text_count}, Görsel: {added_img_count})")

        except Exception:
            pass
